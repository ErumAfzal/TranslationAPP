import streamlit as st
from openai import OpenAI
import fitz
from docx import Document
import io

# ----------------------------
# PAGE CONFIG
# ----------------------------

st.set_page_config(
    page_title="Document Translator",
    page_icon="🌍",
    layout="wide"
)

st.title("🌍 Academic Document Translator")
st.write("Upload PDF, DOCX or TXT and translate into English and/or German.")

# ----------------------------
# API
# ----------------------------

client = OpenAI(
    api_key=st.secrets["API_KEY"],
    base_url="https://api.hrz.uni-giessen.de"
)

# ----------------------------
# SETTINGS
# ----------------------------

with st.sidebar:
    st.header("Settings")

    model_name = st.selectbox(
        "Choose model",
        [
            "openai/gpt-5.5",
            "openai/gpt-5.4"
        ]
    )

    target = st.selectbox(
        "Translate to",
        [
            "English",
            "German",
            "Both English and German"
        ]
    )

# ----------------------------
# FILE EXTRACTORS
# ----------------------------

def extract_pdf(uploaded_file):
    pdf = fitz.open(
        stream=uploaded_file.read(),
        filetype="pdf"
    )

    text = ""

    for page_num, page in enumerate(pdf, start=1):
        text += f"\n\n--- Page {page_num} ---\n\n"
        text += page.get_text("text")
        text += "\n\n"

    return text


def extract_docx(uploaded_file):
    doc = Document(uploaded_file)
    text = []

    for para in doc.paragraphs:
        text.append(para.text)

    return "\n".join(text)


def extract_txt(uploaded_file):
    return uploaded_file.read().decode("utf-8", errors="ignore")


# ----------------------------
# CHUNKING
# ----------------------------

def split_text(text, chunk_size=9000):
    chunks = []

    for i in range(0, len(text), chunk_size):
        chunks.append(text[i:i + chunk_size])

    return chunks


# ----------------------------
# TRANSLATION
# ----------------------------

def translate_chunk(chunk, language, chunk_number, total_chunks):

    prompt = f"""
You are a careful professional academic translator in higher education research.

Translate the following document chunk into {language}.

This is chunk {chunk_number} of {total_chunks}.

STRICT RULES TO AVOID HALLUCINATION:

1. Translate ONLY the text provided in this chunk.
2. Do NOT add new information.
3. Do NOT summarize.
4. Do NOT omit content.
5. Do NOT invent missing words, missing references, missing authors, or missing table values.
6. If any word, sentence, table cell, or reference is unclear because of PDF extraction problems, write [unclear] instead of guessing.
7. Preserve page markers such as --- Page 1 ---.
8. Preserve headings and subheadings.
9. Preserve paragraph order.
10. Preserve citations, footnotes, references, author names, dates, journal names, volume numbers, issue numbers, page numbers, and DOI information.
11. Preserve tables using Markdown table format.
12. Keep all table columns and rows.
13. Do not merge table columns.
14. If the document already contains an official English title, abstract, or keywords, keep the official wording instead of re-translating it.
15. Use natural academic language, but stay faithful to the original.
16. Do not translate Chinese academic expressions too literally when a standard academic English or German equivalent exists.
17. Keep technical terms consistent throughout the chunk.
18. If spacing is broken by PDF extraction, repair spacing only when the correction is obvious.

DOCUMENT CHUNK:

{chunk}
"""

    response = client.chat.completions.create(
        model=model_name,
        messages=[
            {
                "role": "system",
                "content": "You are a precise academic translator. You must not hallucinate or add information."
            },
            {
                "role": "user",
                "content": prompt
            }
        ]
    )

    return response.choices[0].message.content


def translate_document(text, language):

    chunks = split_text(text)
    translated_chunks = []

    progress_bar = st.progress(0)

    total_chunks = len(chunks)

    for i, chunk in enumerate(chunks, start=1):

        translated = translate_chunk(
            chunk,
            language,
            i,
            total_chunks
        )

        translated_chunks.append(translated)

        progress_bar.progress(i / total_chunks)

    return "\n\n".join(translated_chunks)


# ----------------------------
# DOWNLOAD FILES
# ----------------------------

def create_txt(text):
    buffer = io.BytesIO()
    buffer.write(text.encode("utf-8"))
    buffer.seek(0)
    return buffer


def create_docx(text, title):
    doc = Document()
    doc.add_heading(title, level=1)

    for line in text.split("\n"):
        if line.strip():
            doc.add_paragraph(line)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ----------------------------
# FILE UPLOAD
# ----------------------------

uploaded_file = st.file_uploader(
    "Upload PDF, DOCX or TXT",
    type=["pdf", "docx", "txt"]
)

if uploaded_file:

    st.success(f"Uploaded file: {uploaded_file.name}")

    if uploaded_file.name.lower().endswith(".pdf"):
        extracted_text = extract_pdf(uploaded_file)

    elif uploaded_file.name.lower().endswith(".docx"):
        extracted_text = extract_docx(uploaded_file)

    else:
        extracted_text = extract_txt(uploaded_file)

    if not extracted_text.strip():
        st.error("No readable text found. This may be a scanned PDF.")
        st.stop()

    st.subheader("Extracted Text Preview")

    st.text_area(
        "Preview",
        extracted_text[:10000],
        height=250
    )

    st.write(f"Characters extracted: {len(extracted_text)}")

    if st.button("Translate Document"):

        with st.spinner("Translating..."):

            try:

                if target == "English":

                    english = translate_document(
                        extracted_text,
                        "English"
                    )

                    st.subheader("English Translation")

                    st.text_area(
                        "English Output",
                        english,
                        height=500
                    )

                    st.download_button(
                        "Download English TXT",
                        create_txt(english),
                        file_name="english_translation.txt",
                        mime="text/plain"
                    )

                    st.download_button(
                        "Download English DOCX",
                        create_docx(english, "English Translation"),
                        file_name="english_translation.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

                elif target == "German":

                    german = translate_document(
                        extracted_text,
                        "German"
                    )

                    st.subheader("German Translation")

                    st.text_area(
                        "German Output",
                        german,
                        height=500
                    )

                    st.download_button(
                        "Download German TXT",
                        create_txt(german),
                        file_name="german_translation.txt",
                        mime="text/plain"
                    )

                    st.download_button(
                        "Download German DOCX",
                        create_docx(german, "German Translation"),
                        file_name="german_translation.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

                else:

                    english = translate_document(
                        extracted_text,
                        "English"
                    )

                    german = translate_document(
                        extracted_text,
                        "German"
                    )

                    col1, col2 = st.columns(2)

                    with col1:
                        st.subheader("English")

                        st.text_area(
                            "English Output",
                            english,
                            height=500
                        )

                        st.download_button(
                            "Download English TXT",
                            create_txt(english),
                            file_name="english_translation.txt",
                            mime="text/plain"
                        )

                        st.download_button(
                            "Download English DOCX",
                            create_docx(english, "English Translation"),
                            file_name="english_translation.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )

                    with col2:
                        st.subheader("German")

                        st.text_area(
                            "German Output",
                            german,
                            height=500
                        )

                        st.download_button(
                            "Download German TXT",
                            create_txt(german),
                            file_name="german_translation.txt",
                            mime="text/plain"
                        )

                        st.download_button(
                            "Download German DOCX",
                            create_docx(german, "German Translation"),
                            file_name="german_translation.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )

            except Exception as e:
                st.error(str(e))
